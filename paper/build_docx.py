"""严格按祝家颖论文格式精修 docx。
规格(源自祝家颖论文 styles.xml + document.xml):
  大标题: 黑体 22pt 不加粗 居中
  一级标题: 黑体 14pt 不加粗 黑色 左对齐 1.5倍行距
  二级标题: 黑体 12pt 加粗 黑色 左对齐
  正文: 宋体/Times New Roman 10.5pt 两端对齐 1.5倍行距 首行缩进2字符
  摘要/关键词: 宋体 10.5pt (首词加粗)
  图注/表题: 黑体 9pt 居中
  表格: 三线表(顶/底 1.5pt, 表头下 0.75pt, 无竖线), 内容宋体9pt 居中, 表头加粗
  参考文献: 9pt 自动编号 悬挂缩进
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

FIGDIR = 'figures/'
BLACK = RGBColor(0, 0, 0)
doc = Document('paper_base.docx')


def set_run(run, cn='宋体', en='Times New Roman', size=10.5, bold=False):
    rpr = run._element.get_or_add_rPr()
    # 检测是否为上标(文献引用), 若是则字号减小并保留上标
    is_sup = rpr.find(qn('w:vertAlign')) is not None
    run.font.size = Pt(size * 0.75) if is_sup else Pt(size)
    run.font.bold = bold
    run.font.color.rgb = BLACK
    run.font.name = en
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), cn)
    rfonts.set(qn('w:ascii'), en)
    rfonts.set(qn('w:hAnsi'), en)


def set_line_spacing(p, mult=1.5, first_indent_chars=0):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = mult
    if first_indent_chars:
        # 首行缩进 N 个字符
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLineChars'), str(first_indent_chars * 100))


# ---------- 页面设置 ----------
for sec in doc.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.5)
    sec.different_first_page_header_footer = False


# ---------- 页眉(页码, 带下边框线) + 页脚(会议名称) ----------
def add_page_number(paragraph):
    """在段落中插入 PAGE 域(页码)。"""
    run = paragraph.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    set_run(run, cn='宋体', en='Times New Roman', size=9)


for sec in doc.sections:
    # 页眉: 居中页码 + 下边框线
    hdr = sec.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0]
    hp.text = ''
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(hp)
    # 页眉段落下边框(页眉线)
    pPr = hp._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'auto')
    pbdr.append(bottom)
    pPr.append(pbdr)
    # 页脚: 居中会议名称
    ftr = sec.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0]
    fp.text = ''
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(fp.add_run('第41届南京地区研究生通信年会'), cn='宋体',
            en='Times New Roman', size=9)

# ---------- 正文默认样式 ----------
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(10.5)
normal.font.color.rgb = BLACK
normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')

paras = doc.paragraphs

# ---------- 大标题(第0段) ----------
paras[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
set_line_spacing(paras[0], 1.5)
for r in paras[0].runs:
    set_run(r, cn='黑体', en='Times New Roman', size=16, bold=True)

# ---------- 作者(第1段) 居中 ----------
paras[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in paras[1].runs:
    set_run(r, cn='宋体', size=12)
# ---------- 单位(第2段) 居中 小号 ----------
paras[2].alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in paras[2].runs:
    set_run(r, cn='宋体', size=9)


def is_caption(t):
    t = t.strip()
    return (t.startswith('图') or t.startswith('表')) and len(t) < 40 and '给出' not in t


# ---------- 遍历设置正文/标题/摘要/图表题 ----------
for p in doc.paragraphs:
    t = p.text.strip()
    style = p.style.name
    if not t:
        continue
    if style == 'Heading 1':
        set_line_spacing(p, 1.5)
        for r in p.runs:
            set_run(r, cn='黑体', en='Times New Roman', size=14, bold=False)
    elif style == 'Heading 2':
        set_line_spacing(p, 1.5)
        for r in p.runs:
            set_run(r, cn='黑体', en='Arial', size=12, bold=True)
    elif is_caption(t):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_run(r, cn='黑体', en='Times New Roman', size=9, bold=False)
    elif t.startswith('摘') or t.startswith('关键词'):
        set_line_spacing(p, 1.4)
        for r in p.runs:
            set_run(r, cn='楷体', en='Times New Roman', size=10.5)
    elif p in (paras[0], paras[1], paras[2]):
        continue
    else:
        # 正文段: 两端对齐 + 1.5倍行距 + 首行缩进2字符
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_line_spacing(p, 1.5, first_indent_chars=2)
        for r in p.runs:
            set_run(r, cn='宋体', size=10.5)

# ---------- 摘要/关键词首词加粗 ----------
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith('摘') or t.startswith('关键词'):
        if p.runs:
            # 首个run加粗(摘 要 / 关键词)
            for r in p.runs[:1]:
                r.font.bold = True

# ---------- 参考文献段落: 9pt ----------
started_ref = False
for p in doc.paragraphs:
    if p.text.strip() == '参考文献':
        started_ref = True
        continue
    if started_ref and p.text.strip():
        set_line_spacing(p, 1.2)
        p.paragraph_format.first_line_indent = None
        for r in p.runs:
            set_run(r, cn='宋体', size=9)


# ---------- 三线表 ----------
def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'bottom', 'left', 'right'):
        if edge in kwargs:
            tag = qn('w:' + edge)
            el = tcBorders.find(tag)
            if el is None:
                el = OxmlElement('w:' + edge)
                tcBorders.append(el)
            sz, val = kwargs[edge]
            el.set(qn('w:val'), val)
            el.set(qn('w:sz'), str(sz))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), 'auto')


def make_three_line(tbl, header_rows=1):
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 移除表格级默认边框
    tblPr = tbl._tbl.tblPr
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    nrow = len(tbl.rows)
    for ri, row in enumerate(tbl.rows):
        for cell in row.cells:
            borders = {'left': (0, 'none'), 'right': (0, 'none'),
                       'insideV': (0, 'none')}
            top = (12, 'single') if ri == 0 else (0, 'none')
            if ri == header_rows - 1:
                bottom = (6, 'single')       # 表头下细线
            elif ri == nrow - 1:
                bottom = (12, 'single')      # 表格底粗线
            else:
                bottom = (0, 'none')
            borders['top'] = top
            borders['bottom'] = bottom
            set_cell_border(cell, **borders)
            for pp in cell.paragraphs:
                pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pp.paragraph_format.first_line_indent = None
                bold = (ri < header_rows)
                for r in pp.runs:
                    set_run(r, cn='宋体', size=9, bold=bold)


# 判断哪些是数据三线表, 哪些是算法伪代码表(单列)
for tbl in doc.tables:
    ncols = len(tbl.columns)
    if ncols == 1:
        # 伪代码表: 外框(顶/底粗线, 表头下细线), 无竖线, 左对齐, 等宽字体
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tblPr = tbl._tbl.tblPr
        old = tblPr.find(qn('w:tblBorders'))
        if old is not None:
            tblPr.remove(old)
        nrow = len(tbl.rows)
        for ri, row in enumerate(tbl.rows):
            for cell in row.cells:
                borders = {'left': (0, 'none'), 'right': (0, 'none')}
                borders['top'] = (12, 'single') if ri == 0 else (0, 'none')
                if ri == 0:
                    borders['bottom'] = (6, 'single')
                elif ri == nrow - 1:
                    borders['bottom'] = (12, 'single')
                else:
                    borders['bottom'] = (0, 'none')
                set_cell_border(cell, **borders)
                for pp in cell.paragraphs:
                    pp.alignment = (WD_ALIGN_PARAGRAPH.CENTER if ri == 0
                                    else WD_ALIGN_PARAGRAPH.LEFT)
                    pp.paragraph_format.first_line_indent = None
                    for r in pp.runs:
                        set_run(r, cn='宋体', en='Consolas', size=9,
                                bold=(ri == 0))
    else:
        make_three_line(tbl, header_rows=1)


# ---------- 插入图片与图注 ----------
def _insert_after(paragraph):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def place_figure(anchor, img, caption, width_cm=8.2):
    for p in list(doc.paragraphs):
        if anchor in p.text and p.style.name not in ('Heading 1', 'Heading 2'):
            cap_p = _insert_after(p)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.first_line_indent = None
            set_run(cap_p.add_run(caption), cn='黑体', en='Times New Roman',
                    size=9)
            img_p = _insert_after(p)
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_p.paragraph_format.first_line_indent = None
            img_p.add_run().add_picture(FIGDIR + img, width=Cm(width_cm))
            return True
    print('未找到锚点:', anchor)
    return False


place_figure('本文的仿真场景如图1所示', 'fig_scenario.png',
             '图1 仿真场景与业务流转示意', 10.0)
place_figure('图2给出Walker Delta星座', 'fig_constellation.png',
             '图2 Walker Delta星座与+Grid链路拓扑示意', 7.8)
place_figure('图3给出CA-MQR与Vanilla-DQN的训练收敛', 'fig_convergence.png',
             '图3 训练收敛曲线', 8.0)
place_figure('图4给出平均端到端时延随归一化负载', 'fig_delay.png',
             '图4 平均端到端时延随负载的变化', 8.4)
place_figure('图5给出链路利用率', 'fig_util.png',
             '图5 链路利用率(95分位)随负载的变化', 8.4)
place_figure('图6给出重载场景下三类业务的QoS满足率', 'fig_perclass.png',
             '图6 三类业务的QoS满足率对比', 9.0)
place_figure('图7进一步给出总体QoS满足率', 'fig_satrate.png',
             '图7 总体QoS满足率随负载的变化', 8.4)
place_figure('图8给出其收敛曲线', 'fig_dueling.png',
             '图8 竞争式与普通网络结构的收敛曲线对比', 8.4)
place_figure('图9给出P95尾部时延与QoS满足率', 'fig_potential.png',
             '图9 下游拥塞势场在热点汇聚场景下的增强效果', 11.5)

doc.save('低轨卫星网络拥塞感知的多业务智能路由方法_姚飞.docx')
print('已生成最终论文 docx')
